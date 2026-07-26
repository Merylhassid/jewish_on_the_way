from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = Path(__file__).with_name("Daniel_Yehudai_Resume.docx")
BLUE = "1F4F85"
TEXT = "172033"
MUTED = "46566A"
LINE = "B8C8D9"


def set_cell_free_spacing(paragraph, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.line_spacing = line


def set_run(run, size=9.4, bold=False, color=TEXT, italic=False):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    return run


def add_hyperlink(paragraph, text, url, size=8.35):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), TEXT)
    properties.append(color)
    size_node = OxmlElement("w:sz")
    size_node.set(qn("w:val"), str(round(size * 2)))
    properties.append(size_node)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    properties.append(fonts)
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_separator(paragraph, text=" | ", size=8.35):
    set_run(paragraph.add_run(text), size=size)


def add_bottom_border(paragraph):
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "5")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), LINE)
    borders.append(bottom)


def add_section_heading(document, title):
    paragraph = document.add_paragraph()
    set_cell_free_spacing(paragraph, before=5.5, after=2.3)
    set_run(paragraph.add_run(title.upper()), size=10.3, bold=True, color=BLUE)
    add_bottom_border(paragraph)
    return paragraph


def add_body(document, text, size=9.15, after=0):
    paragraph = document.add_paragraph()
    set_cell_free_spacing(paragraph, after=after, line=1.0)
    set_run(paragraph.add_run(text), size=size)
    return paragraph


def add_label_line(document, label, text, size=8.85):
    paragraph = document.add_paragraph()
    set_cell_free_spacing(paragraph, after=0)
    set_run(paragraph.add_run(label), size=size, bold=True, color="233D5D")
    set_run(paragraph.add_run(text), size=size)
    return paragraph


def add_entry_heading(document, title, date=None):
    paragraph = document.add_paragraph()
    set_cell_free_spacing(paragraph, before=0.5, after=0.2)
    if date:
        paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(17.6), WD_TAB_ALIGNMENT.RIGHT)
    set_run(paragraph.add_run(title), size=9.35, bold=True)
    if date:
        set_run(paragraph.add_run("\t" + date), size=9.1, bold=True, color=MUTED)
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    fmt = paragraph.paragraph_format
    fmt.left_indent = Cm(0.45)
    fmt.first_line_indent = Cm(-0.25)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0.7)
    fmt.line_spacing = 1.0
    set_run(paragraph.add_run(text), size=8.85)
    return paragraph


document = Document()
section = document.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.05)
section.bottom_margin = Cm(1.0)
section.left_margin = Cm(1.4)
section.right_margin = Cm(1.4)
section.header_distance = Cm(0.3)
section.footer_distance = Cm(0.3)

styles = document.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(9.2)
normal.font.color.rgb = RGBColor.from_string(TEXT)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

title = document.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_cell_free_spacing(title, after=0)
set_run(title.add_run("Daniel Yehudai"), size=21.5, bold=True)

subtitle = document.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_cell_free_spacing(subtitle, after=2)
set_run(
    subtitle.add_run("Software Engineering Graduate | Full-Stack Developer"),
    size=11.2,
    bold=True,
    color=BLUE,
)

contact = document.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_cell_free_spacing(contact, after=4)
add_hyperlink(contact, "050-821-1682", "tel:+972508211682")
add_separator(contact)
add_hyperlink(contact, "daniyehudai@gmail.com", "mailto:daniyehudai@gmail.com")
add_separator(contact)
add_hyperlink(
    contact,
    "linkedin.com/in/daniel-yehudai-30a454318",
    "https://www.linkedin.com/in/daniel-yehudai-30a454318",
)
add_separator(contact)
add_hyperlink(contact, "github.com/Daniely2090", "https://github.com/Daniely2090")
add_bottom_border(contact)

add_section_heading(document, "Profile")
add_body(
    document,
    "Software Engineering graduate with hands-on experience developing a full-stack mobile application using React Native, Expo, NestJS, TypeScript, and PostgreSQL/PostGIS. Built REST APIs, geospatial search, JWT-based authentication, real-time communication, and an explainable NLP search pipeline. Motivated by turning product requirements into reliable, user-focused software.",
    size=8.95,
)

add_section_heading(document, "Technical Skills")
add_label_line(document, "Languages: ", "TypeScript, JavaScript, Python | Academic familiarity: Java, C++")
add_label_line(document, "Frontend & Mobile: ", "React Native, Expo, Expo Router")
add_label_line(document, "Backend: ", "NestJS, Node.js, REST APIs, JWT authentication, Socket.IO")
add_label_line(document, "Data: ", "PostgreSQL, PostGIS, SQL, TypeORM")
add_label_line(document, "Machine Learning: ", "Text classification, TF-IDF, Naive Bayes, data preprocessing, train/validation/test evaluation, Precision, Recall, F1", size=8.65)
add_label_line(document, "Tools: ", "Git, GitHub, Jira, Swagger, PM2")

add_section_heading(document, "Project")
add_entry_heading(document, "Jewish On The Way | Final-Year Team Project")
technology = document.add_paragraph()
set_cell_free_spacing(technology, after=1)
set_run(
    technology.add_run("React Native, Expo, NestJS, TypeScript, PostgreSQL, PostGIS, TypeORM, Socket.IO"),
    size=8.65,
    color=MUTED,
    italic=True,
)
add_bullet(document, "Developed a cross-platform mobile application that helps Jewish travelers find kosher restaurants, synagogues, minyans, and Shabbat hosting by destination or current location.")
add_bullet(document, "Built an explainable multilingual search pipeline using TF-IDF and Naive Bayes classifiers, with Python-based training and TypeScript inference; achieved 82.6% accuracy on a held-out test set.")
add_bullet(document, "Designed REST APIs and relational data flows with NestJS, TypeORM, and PostgreSQL; implemented JWT authentication with refresh tokens and role-protected operations.")
add_bullet(document, "Implemented location-based discovery with PostGIS distance queries, plus real-time chat and hosting updates using Socket.IO.")
add_bullet(document, "Deployed the backend to a Linux VPS with PM2 and a cloud-hosted PostgreSQL database; documented APIs with Swagger and supported the backend with automated tests.")

add_section_heading(document, "Education")
add_entry_heading(document, "B.Sc. in Software Engineering | SCE – Shamoon College of Engineering", "2021–2026")
add_body(document, "Moshal Program Scholar — Selected for an international excellence and leadership program for high-achieving students.", size=8.85)

add_section_heading(document, "Military Experience")
add_entry_heading(document, "Planning and Production Control Technician | Israel Defense Forces", "2017–2019")
add_body(document, "Coordinated operational workflows, maintained schedules and documentation, monitored task progress, and supported on-time execution in a structured, team-based environment.", size=8.85)

add_section_heading(document, "Volunteer Experience")
add_entry_heading(document, "Matityahu Elementary School", "2018–Present")
add_body(document, "Provide educational and social support to children in special education.", size=8.8)
add_entry_heading(document, "Tzamid Program", "2021–2022")
add_body(document, "Supported social and educational activities for children with disabilities.", size=8.8)

add_section_heading(document, "Languages")
languages = document.add_paragraph()
set_cell_free_spacing(languages)
set_run(languages.add_run("Hebrew: "), size=8.9, bold=True)
set_run(languages.add_run("Native  |  "), size=8.9)
set_run(languages.add_run("English: "), size=8.9, bold=True)
set_run(languages.add_run("Working proficiency"), size=8.9)

properties = document.core_properties
properties.title = "Daniel Yehudai Resume"
properties.subject = "Software Engineering Graduate | Full-Stack Developer"
properties.author = "Daniel Yehudai"
properties.keywords = "Software Engineering, Full-Stack, NestJS, React Native, TypeScript"

document.save(OUTPUT)
print(OUTPUT)
